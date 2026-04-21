from dataclasses import dataclass
from typing import List
import os
from docx import Document
from PyPDF2 import PdfReader


@dataclass
class ExtractedDocument:
    text: str
    tables: List[List[List[str]]]


class TextExtractor:

    SUPPORTED_EXTENSIONS = (".txt", ".pdf", ".docx", ".doc")

    def extract(self, file_path: str) -> ExtractedDocument:

        if not os.path.exists(file_path):
            raise FileNotFoundError(file_path)

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".txt":
            return ExtractedDocument(
                text=self._extract_txt(file_path),
                tables=[]
            )

        if ext == ".pdf":
            return ExtractedDocument(
                text=self._extract_pdf(file_path),
                tables=[]
            )

        return self._extract_docx(file_path)

    def _extract_txt(self, path: str) -> str:
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _extract_pdf(self, path: str) -> str:
        reader = PdfReader(path)
        text = ""
        for p in reader.pages:
            t = p.extract_text()
            if t:
                text += t + "\n"
        return text

    def _extract_docx(self, path: str) -> ExtractedDocument:
        doc = Document(path)

        full_content = []

        for p in doc.paragraphs:
            if p.text.strip():
                full_content.append(p.text.strip())

        tables_data = []
        for table in doc.tables:
            t_rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    c_text = cell.text.strip()
                    if not cells or c_text != cells[-1]:
                        cells.append(c_text)

                if any(cells):
                    t_rows.append(cells)
                    full_content.append(" | ".join([c for c in cells if c]))

            tables_data.append(t_rows)

        return ExtractedDocument(
            text="\n".join(full_content),
            tables=tables_data
        )